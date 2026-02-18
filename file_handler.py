import fitz          # PyMuPDF
import docx
import os
import re
from werkzeug.utils import secure_filename

ALLOWED_EXTENSIONS = {"pdf", "docx", "doc"}

def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_pdf(filepath: str) -> str:
    """Ekstrak teks dari file PDF menggunakan PyMuPDF."""
    text_parts = []
    with fitz.open(filepath) as doc:
        for page in doc:
            text_parts.append(page.get_text("text"))
    return "\n".join(text_parts)


def extract_text_from_docx(filepath: str) -> str:
    """Ekstrak teks dari file DOCX."""
    document = docx.Document(filepath)
    paragraphs = [para.text for para in document.paragraphs if para.text.strip()]
    
    # Ekstrak juga teks dari tabel
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    
    return "\n".join(paragraphs)


def extract_text(filepath: str, file_type: str) -> str:
    """Router utama ekstraksi teks."""
    if file_type == "pdf":
        return extract_text_from_pdf(filepath)
    elif file_type in ("docx", "doc"):
        return extract_text_from_docx(filepath)
    else:
        raise ValueError(f"Tipe file tidak didukung: {file_type}")


def clean_text(text: str) -> str:
    """Bersihkan teks dari karakter tidak perlu."""
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[^\w\s.,;:!?\"'()\-–—@#&%/\\]", "", text)
    return text.strip()


def save_upload(file, upload_folder: str) -> tuple[str, str]:
    """Simpan file upload dan kembalikan (filepath, file_type)."""
    os.makedirs(upload_folder, exist_ok=True)
    filename  = secure_filename(file.filename)
    file_type = filename.rsplit(".", 1)[1].lower()
    filepath  = os.path.join(upload_folder, filename)
    file.save(filepath)
    return filepath, file_type