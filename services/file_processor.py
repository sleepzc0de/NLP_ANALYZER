import os
import fitz  # PyMuPDF
from docx import Document as DocxDocument


class FileProcessor:

    @staticmethod
    def extract_text_from_pdf(filepath: str) -> str:
        text_parts = []
        try:
            doc = fitz.open(filepath)
            if doc.is_encrypted:
                raise RuntimeError(
                    "File PDF terproteksi password, tidak bisa diekstrak."
                )
            for page_num, page in enumerate(doc):
                page_text = page.get_text("text")
                if page_text.strip():
                    text_parts.append(page_text)
            doc.close()
        except RuntimeError:
            raise
        except Exception as e:
            raise RuntimeError(f"Gagal membaca PDF: {e}")

        result = "\n".join(text_parts).strip()
        if not result:
            raise RuntimeError(
                "PDF tidak mengandung teks yang bisa diekstrak. "
                "Kemungkinan PDF berisi gambar/scan saja."
            )
        return result

    @staticmethod
    def extract_text_from_docx(filepath: str) -> str:
        try:
            doc = DocxDocument(filepath)
            paragraphs = []

            # Ambil teks dari paragraf utama
            for p in doc.paragraphs:
                if p.text.strip():
                    paragraphs.append(p.text.strip())

            # Ambil teks dari tabel
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())

            result = "\n".join(paragraphs).strip()
            if not result:
                raise RuntimeError(
                    "Dokumen DOCX tidak mengandung teks yang bisa diekstrak."
                )
            return result

        except RuntimeError:
            raise
        except Exception as e:
            raise RuntimeError(f"Gagal membaca DOCX: {e}")

    @classmethod
    def extract_text(cls, filepath: str, file_ext: str) -> str:
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File tidak ditemukan: {filepath}")

        ext = file_ext.lower().lstrip(".")
        if ext == "pdf":
            return cls.extract_text_from_pdf(filepath)
        elif ext in ("docx", "doc"):
            return cls.extract_text_from_docx(filepath)
        else:
            raise ValueError(f"Format tidak didukung: {ext}")